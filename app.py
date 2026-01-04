import os
import streamlit as st
import google.generativeai as genai
from PIL import Image
from dotenv import load_dotenv
import database  # Pastikan file database.py ada di folder yang sama
from docx import Document
from fpdf import FPDF
from io import BytesIO

# --- 1. SETUP KONFIGURASI & THEME ---
load_dotenv()
api_key = os.getenv("GEMINI_API_KEY")

if not api_key:
    st.error("API Key tidak ditemukan!")
    st.stop()

genai.configure(api_key=api_key)
model = genai.GenerativeModel('gemini-flash-latest')

st.set_page_config(page_title="InsightVision", layout="wide", page_icon="📸")

# --- 2. FUNGSI EKSPOR (DOCX & PDF) ---
def buat_docx(teks):
    doc = Document()
    doc.add_heading('Laporan Analisis InsightVision', 0)
    doc.add_paragraph(teks)
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()

def buat_pdf(teks):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    # Membersihkan karakter non-latin untuk menghindari error library FPDF standar
    clean_text = teks.encode('latin-1', 'ignore').decode('latin-1')
    pdf.multi_cell(0, 10, txt=clean_text)
    return pdf.output(dest='S').encode('latin-1')

# --- 3. CSS CUSTOM (TAMPILAN MEWAH) ---
st.markdown("""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;600&display=swap');
    html, body, [class*="css"]  {
        font-family: 'Inter', sans-serif;
    }

    /* Efek Glassmorphism pada Sidebar */
    [data-testid="stSidebar"] {
        background: rgba(28, 31, 46, 0.95);
        border-right: 1px solid rgba(255, 255, 255, 0.1);
    }

    /* Styling Button Utama */
    .stButton>button {
        width: 100%;
        border-radius: 12px;
        height: 3em;
        background-image: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
        border: none;
        font-weight: 600;
        transition: all 0.3s ease;
    }
    .stButton>button:hover {
        transform: translateY(-2px);
        box-shadow: 0 5px 15px rgba(0,0,0,0.3);
        color: white;
    }

    /* Download Buttons Styling */
    div.stDownloadButton > button {
        background-color: #2e3b4e !important;
        color: white !important;
        border: 1px solid rgba(255,255,255,0.2) !important;
    }
    
    .footer {
        position: fixed;
        bottom: 10px;
        text-align: center;
        width: 100%;
        color: gray;
        font-size: 12px;
    }
    </style>
    """, unsafe_allow_html=True)

# --- 4. SIDEBAR RIWAYAT ---
with st.sidebar:
    st.image("https://cdn-icons-png.flaticon.com/512/1055/1055666.png", width=80)
    st.title("InsightVision")
    st.markdown("---")
    st.subheader("📜 Riwayat Terakhir")
    riwayat = database.ambil_riwayat()
    if riwayat:
        for data in riwayat:
            with st.expander(f"📄 {data['nama_file']}"):
                st.caption(f"📅 {data['waktu_simpan']}")
                st.write(data['hasil_analisis'][:100] + "...")
    else:
        st.info("Belum ada riwayat.")

# --- 5. LOGIKA RESET ---
if "file_sebelumnya" not in st.session_state:
    st.session_state["file_sebelumnya"] = None

st.title("📸 InsightVision")
st.markdown("<p style='color: gray; font-size: 18px;'>Transformasi Gambar Menjadi Wawasan Cerdas</p>", unsafe_allow_html=True)

uploaded_file = st.file_uploader("Seret dan lepaskan gambar di sini", type=["jpg", "png", "jpeg"])

if uploaded_file is not None:
    if st.session_state["file_sebelumnya"] != uploaded_file.name:
        if 'hasil_terakhir' in st.session_state:
            del st.session_state['hasil_terakhir']
        st.session_state["file_sebelumnya"] = uploaded_file.name
        st.rerun()
else:
    if st.session_state["file_sebelumnya"] is not None:
        if 'hasil_terakhir' in st.session_state:
            del st.session_state['hasil_terakhir']
        st.session_state["file_sebelumnya"] = None
        st.rerun()

# --- 6. TAMPILAN UTAMA (KOLOM) ---
col1, col2 = st.columns([1, 1.2], gap="large")

with col1:
    if uploaded_file is not None:
        image = Image.open(uploaded_file)
        st.image(image, caption="Preview Gambar", use_container_width=True)
        
        btn_analisis = st.button("🚀 Mulai Analisis Cerdas")
        
        if btn_analisis:
            with st.spinner("Menganalisis konten gambar..."):
                try:
                    prompt = """
                    Kamu adalah InsightVision. Berikan laporan analisis gambar yang sangat mendalam:
                    1. **📊 Ringkasan**: Deskripsi tingkat tinggi.
                    2. **🔍 Detail Visual**: Objek, warna, komposisi.
                    3. **📝 Teks/OCR**: Ekstraksi semua tulisan yang terlihat.
                    4. **💡 Interpretasi**: Konteks dan saran pakar praktis.
                    Gunakan Bahasa Indonesia profesional.
                    """
                    response = model.generate_content([prompt, image])
                    st.session_state['hasil_terakhir'] = response.text
                    database.simpan_ke_db(uploaded_file.name, response.text)
                except Exception as e:
                    st.error(f"Error: {e}")

with col2:
    if 'hasil_terakhir' in st.session_state:
        st.success("Analisis Berhasil!")
        st.markdown(st.session_state['hasil_terakhir'])
        
        # --- FITUR EKSPOR (DOCX & PDF) ---
        st.markdown("---")
        st.subheader("📥 Ekspor Laporan")
        
        btn_col1, btn_col2 = st.columns(2)
        
        with btn_col1:
            docx_data = buat_docx(st.session_state['hasil_terakhir'])
            st.download_button(
                label="📄 Download DOCX",
                data=docx_data,
                file_name=f"Insight_{uploaded_file.name}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            
        with btn_col2:
            try:
                pdf_data = buat_pdf(st.session_state['hasil_terakhir'])
                st.download_button(
                    label="📕 Download PDF",
                    data=pdf_data,
                    file_name=f"Insight_{uploaded_file.name}.pdf",
                    mime="application/pdf"
                )
            except Exception as e:
                st.error("Karakter teks tidak didukung untuk PDF standar.")

# --- 7. CHAT MULTIMODAL ---
if uploaded_file is not None and 'hasil_terakhir' in st.session_state:
    st.markdown("---")
    st.subheader("💬 Diskusi Lanjutan")
    user_question = st.text_input("Tanyakan detail spesifik tentang gambar ini:", key=f"chat_{uploaded_file.name}", placeholder="Contoh: Apa arti teks di pojok kanan?")

    if user_question:
        with st.chat_message("assistant"):
            with st.spinner("Mengkaji ulang..."):
                try:
                    chat_res = model.generate_content([
                        f"Konteks: {st.session_state['hasil_terakhir']}\n\nPertanyaan: {user_question}", 
                        image
                    ])
                    st.write(chat_res.text)
                except Exception as e:
                    st.error(f"Gagal merespon: {e}")

st.markdown("<br><br><div class='footer'>InsightVision v1.1 | 2026</div>", unsafe_allow_html=True)
